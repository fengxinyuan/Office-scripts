# Desc: 根据excel批量生成word文档
#       1. 读取excel文件中的数据
#       2. 复制模板并重命名
#       3. 替换文件中的关键字
#       4. 保存文件
# TIPS：
#       1. 为了保留原格式，使用了docx库的文本块.runs，遍历paragraph.runs使用run.text.replace替换关键字
#           不使用paragraph.text.replace直接替换，否则会丢失原格式
#       2. 需要检查文本块分割是否正确，否则无法替换关键字
#       3. word和excel里写出关键字，并且关键字一一对应，存在key列表中

from docx import Document
from shutil import copy
from pandas import read_excel
import os

# 关键词
keys = [
    "序号",
    "姓名",
    "性别",
    "手机号",
    "入学年份",
    "学院",
    "班级",
    "学号",
    "备注1",
    "当前学期",
    "选修课程",
    "课程讲师",
    "教室",
    "备注2",
]

model_path = "模板.doc"
excel_path = "名单.xlsx"
file_keyword = "姓名"
dir_name = "files"
file_names = []

def test():
    print("\033[91m测试开始\n\033[0m")
    doc = Document(model_path)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            print(run.text)
    print("\033[91m\n测试成功，检查文本块分割是否正确\n\033[0m")


def readfile():
    # 打开excel文件并获取关键列的数据
    contents = read_excel(excel_path)
    file_names.extend(contents[file_keyword].tolist())
    print("\033[91m读取成功，开始生成\n\033[0m")


def makefile():
    # 打开excel文件
    contents = read_excel(excel_path)
    # 创建文件夹
    if not os.path.exists(dir_name):
        os.mkdir(dir_name)
    # 进入文件夹
    os.chdir(dir_name)

    for i, name in enumerate(file_names):
        # 复制模板并重命名
        copy(f"../{model_path}", f"{name}.doc")
        doc = Document(f"{name}.doc")

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                for key in keys:
                    if key in run.text:
                        run.text = run.text.replace(key, str(contents[key][i]))

        # 保存文档
        doc.save(f"{name}.doc")
        print(f"{name}.doc 生成成功")

    print("\033[91m全部生成成功\033[0m")


if __name__ == "__main__":
    test()
    readfile()
    makefile()
